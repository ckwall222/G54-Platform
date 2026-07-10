"""
Generate SOW Word document for Great Mountain West
Run: python scripts/generate_sow.py
Output: C:\Projects\GMW\docs\sow\SOW_GreatMountainWest.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0F, 0x17, 0x2A)   # headings
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)   # accent / table header
LIGHT_BLUE  = RGBColor(0xDB, 0xEA, 0xFE)   # table header fill
MID_GRAY    = RGBColor(0x47, 0x55, 0x69)   # body text
LIGHT_GRAY  = RGBColor(0xF1, 0xF5, 0xF9)   # alt row fill
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
RED         = RGBColor(0xDC, 0x26, 0x26)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)   # RGBColor.__str__ returns 6-char hex e.g. '1D4ED8'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_type='bottom', color='E2E8F0', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{border_type}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:color'), color)
    tcBorders.append(border)

def add_page_break(doc):
    doc.add_page_break()

def set_run_font(run, size=None, bold=False, italic=False, color=None, name='Calibri'):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, color=DARK_NAVY, size=None, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.name = 'Calibri'
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = color
    if level == 1:
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1D4ED8')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, color=MID_GRAY, size=10.5, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p

def add_bullet(doc, text, color=MID_GRAY, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p

def add_numbered(doc, text, color=MID_GRAY, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p

def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, ACCENT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        set_run_font(run, size=9.5, bold=True, color=WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = LIGHT_GRAY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if isinstance(val, tuple):
                text, bold, color = val
            else:
                text, bold, color = str(val), False, MID_GRAY
            run = p.add_run(text)
            set_run_font(run, size=9.5, bold=bold, color=color)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Confidential — Great Mountain West / Christopher Wall   |   Page ')
    set_run_font(run, size=8, color=MID_GRAY)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = p.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    set_run_font(run2, size=8, color=MID_GRAY)

def set_margins(doc, top=1.0, bottom=1.0, left=1.15, right=1.15):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

# ── Cover / Header block ──────────────────────────────────────────────────────

def build_cover(doc):
    # Top accent bar (simulated with a table)
    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = bar.rows[0].cells[0]
    set_cell_bg(cell, ACCENT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('G54 PLATFORM — ENTERPRISE RESOURCE PLATFORM')
    set_run_font(run, size=11, bold=True, color=WHITE)
    cell.width = Inches(6.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Title block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Statement of Work')
    set_run_font(run, size=28, bold=True, color=DARK_NAVY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run('Sage 100 Replacement — Full ERP, CRM & Commerce Platform')
    set_run_font(run2, size=13, color=MID_GRAY, italic=True)

    add_divider(doc)

    # Metadata table
    meta = doc.add_table(rows=8, cols=2)
    meta.style = 'Table Grid'
    fields = [
        ('Client',              'Great Mountain West'),
        ('Client Website',      'g54.com'),
        ('Client Contact',      '___________________________'),
        ('Service Provider',    'Christopher Wall'),
        ('Provider Email',      'ck.wall@icloud.com'),
        ('Document Version',    '1.0 — Draft'),
        ('Date',                'July 10, 2026'),
        ('Document Status',     'Pending Client Review & Execution'),
    ]
    for i, (label, value) in enumerate(fields):
        row = meta.rows[i]
        bg = LIGHT_GRAY if i % 2 == 0 else WHITE
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after = Pt(5)
        lr = lp.add_run(label)
        set_run_font(lr, size=10, bold=True, color=DARK_NAVY)
        vp = row.cells[1].paragraphs[0]
        vp.paragraph_format.space_before = Pt(5)
        vp.paragraph_format.space_after = Pt(5)
        vr = vp.add_run(value)
        set_run_font(vr, size=10, color=MID_GRAY)
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.3)

    add_page_break(doc)

# ── Document body ─────────────────────────────────────────────────────────────

def build_body(doc):

    # 1. Background
    add_heading(doc, '1.  Project Background & Overview', 1)
    add_body(doc, (
        'Great Mountain West currently operates on Sage 100 as its primary business management platform. '
        'Following an evaluation of Odoo as a potential replacement — which did not satisfy the depth of '
        'financial functionality required — Great Mountain West is commissioning a purpose-built, '
        'cloud-native enterprise platform.'
    ))
    add_body(doc, 'This platform will consolidate and replace the following operational domains:')
    for item in [
        'Financial management and accounting (replacing Sage 100)',
        'Customer Relationship Management (CRM)',
        'eCommerce — customer-facing order entry and management',
        'Inventory and warehouse management',
        'Point of Sale (POS)',
        'Project and job management',
        'Accounts Receivable and Accounts Payable workflows',
    ]:
        add_bullet(doc, item)

    add_body(doc, (
        'The resulting system will be a single, integrated, cloud-hosted platform accessible via internet '
        'browser, designed to meet PCI DSS and SOC 2 Type 2 compliance requirements from the ground up.'
    ), space_after=4)

    # 2. Objectives
    add_heading(doc, '2.  Project Objectives', 1)
    for item in [
        'Deliver a full-replacement ERP platform purpose-built for Great Mountain West\'s operational model.',
        'Eliminate dependency on Sage 100 and any associated legacy integrations.',
        'Provide a secure, compliance-ready platform with end-to-end data security.',
        'Enable role-based access across all departments.',
        'Build iteratively using Agile sprint methodology, delivering working software at each phase boundary.',
        'Architect for scalability, backup integrity, and disaster recovery from day one.',
    ]:
        add_bullet(doc, item)

    # 3. Phase overview table
    add_heading(doc, '3.  Project Phases Overview', 1)
    styled_table(doc,
        headers=['Phase', 'Name', 'Methodology', 'Pricing Model'],
        rows=[
            [('Phase 1', True, DARK_NAVY), 'Discovery, Architecture & Design', 'Flat-rate delivery', ('$5,000', True, ACCENT_BLUE)],
            [('Phase 2', True, DARK_NAVY), 'Core Platform Foundation',         'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 3', True, DARK_NAVY), 'Financial Suite (GL, AR, AP)',     'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 4', True, DARK_NAVY), 'CRM & Sales',                      'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 5', True, DARK_NAVY), 'Inventory, POS & eCommerce',       'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 6', True, DARK_NAVY), 'Project & Job Management',         'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 7', True, DARK_NAVY), 'Reporting, Analytics & Dashboards','Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
            [('Phase 8', True, DARK_NAVY), 'Compliance, Security & Launch',    'Sprint-based',       ('TBD post-Phase 1', False, MID_GRAY)],
        ],
        col_widths=[0.65, 2.5, 1.5, 1.65],
    )
    add_body(doc, (
        'Note: Phases 2–8 will be scoped, estimated, and priced as fixed deliverables following '
        'the completion and approval of Phase 1. Sprint pricing will be formalized in a Phase 2+ amendment.'
    ), color=MID_GRAY, size=9.5)

    # 4. Phase 1 detail
    add_heading(doc, '4.  Phase 1 — Discovery, Architecture & Design', 1)

    add_heading(doc, '4.1  Scope', 2)
    add_body(doc, (
        'Phase 1 is a flat-rate discovery and design engagement. Its purpose is to produce the complete '
        'blueprint — functional, technical, and visual — that will govern all subsequent development phases.'
    ))

    add_heading(doc, '4.1.1  Discovery & Requirements Workshops', 3)
    for item in [
        'Structured interviews and working sessions with stakeholders across all departments',
        'Documentation of current Sage 100 workflows and data models',
        'Gap analysis between current state and desired future state',
        'Identification of third-party integrations (payment processors, shipping, email, file storage)',
        'Compliance scoping session (PCI DSS, SOC 2 Type 2)',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.1.2  Persona Development', 3)
    add_body(doc, 'Definition and documentation of all user personas, including:')
    for item in [
        'Business personas: Executive, Finance Manager, Sales Rep, Sales Manager, Art Director, '
        'Graphic Designer, Prepress Technician, Production Manager, Production Operator, AR Clerk, '
        'AP Clerk, Customer Service Rep, Warehouse Manager, Project Manager, IT Administrator',
        'System personas: Super Admin, API Consumer, Payment Gateway, Email Service, '
        'Reporting Engine, Backup Service, Audit Engine, CI/CD Pipeline, Session Manager',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.1.3  Full Epic & User Story Documentation', 3)
    for item in [
        'Complete product backlog organized by module and persona',
        'Acceptance criteria for each user story',
        'Priority and dependency mapping',
        'Estimated story point ranges for sprint planning reference',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.1.4  Feature Requirements Specification', 3)
    for item in [
        'Functional requirements document (FRD) covering all modules',
        'Non-functional requirements (performance, security, availability, compliance)',
        'Data model and entity relationship diagrams (ERD)',
        'Role-based access control (RBAC) matrix',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.1.5  UX/UI Design & Wireframes', 3)
    for item in [
        'Information architecture and site map',
        'Low-fidelity wireframes for all primary user flows',
        'High-fidelity mockups for key screens',
        'Interactive prototype for client review and approval',
        'Design system foundations (typography, color, component library)',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.1.6  Technical Architecture Design', 3)
    for item in [
        'Platform and infrastructure selection',
        'Security architecture (authentication, authorization, encryption)',
        'Backup and disaster recovery architecture',
        'CI/CD pipeline architecture',
        'PCI DSS and SOC 2 Type 2 control mapping',
        'Third-party service selection and integration design',
        'API architecture design',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.2  Phase 1 Deliverables', 2)
    styled_table(doc,
        headers=['#', 'Deliverable', 'Format'],
        rows=[
            ['1',  'Discovery Workshop Summary & Current State Analysis',   'PDF / Notion'],
            ['2',  'Persona Documentation (business and system personas)',   'PDF / Notion'],
            ['3',  'Complete Product Backlog (Epics, Stories, Acceptance Criteria)', 'PDF / Notion'],
            ['4',  'Functional Requirements Document (FRD)',                'PDF'],
            ['5',  'Non-Functional Requirements & Compliance Mapping',      'PDF'],
            ['6',  'Data Model & Entity Relationship Diagrams (ERD)',       'PDF / Diagram'],
            ['7',  'RBAC Matrix',                                           'PDF / Spreadsheet'],
            ['8',  'UX Wireframes — all primary flows',                     'Figma'],
            ['9',  'High-Fidelity UI Mockups — key screens',               'Figma'],
            ['10', 'Interactive Prototype',                                 'Figma (shareable link)'],
            ['11', 'Technical Architecture Document',                       'PDF'],
            ['12', 'Phase 2–8 Refined Scope, Sprint Plan & Cost Estimate', 'PDF'],
        ],
        col_widths=[0.3, 4.0, 2.0],
    )

    add_heading(doc, '4.3  Investment', 2)
    styled_table(doc,
        headers=['Item', 'Detail'],
        rows=[
            [('Phase 1 Total', True, DARK_NAVY),     ('$5,000.00 USD', True, ACCENT_BLUE)],
            ['Pricing Model',                          'Flat-rate delivery (not time-bound)'],
            ['Payment Terms',                          '100% due at project start, prior to commencement'],
            ['Accepted Payment Methods',               'ACH, Wire Transfer, Check, Zelle'],
        ],
        col_widths=[2.3, 4.0],
    )

    add_heading(doc, '4.4  Revisions', 2)
    add_body(doc, (
        'Phase 1 includes up to two (2) rounds of revisions on wireframes and mockups following initial '
        'delivery. Additional revision rounds may be requested at a rate of $175.00 per hour.'
    ))

    add_heading(doc, '4.5  Client Responsibilities', 2)
    for item in [
        'Provide access to at least one designated stakeholder per department for discovery sessions',
        'Provide read-only access to Sage 100 data and report samples as needed',
        'Respond to document review requests within five (5) business days',
        'Consolidate internal feedback before submitting to the project team',
        'Provide existing brand assets (logo, style guide, brand colors)',
    ]:
        add_bullet(doc, item)

    add_page_break(doc)

    # 5. Phase 2-8 overview
    add_heading(doc, '5.  Phase 2–8 — Development Overview', 1)
    add_body(doc, (
        'Development phases will be executed using Agile sprint methodology, with each sprint '
        'delivering working, testable software. Sprint details will be formalized in a Phase 2+ '
        'SOW amendment following client approval of Phase 1 deliverables.'
    ))

    add_heading(doc, '5.1  Sprint Model', 2)
    for item in [
        'Sprint duration: 2 weeks',
        'Each sprint includes: planning, development, testing, client demo, retrospective',
        'All sprints will include automated test coverage',
        'A staging environment will be maintained throughout development',
        'Production deployments will follow sprint acceptance by the client',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '5.2  Module Development Roadmap', 2)
    styled_table(doc,
        headers=['Phase', 'Sprint Group', 'Modules & Deliverables'],
        rows=[
            ['Phase 2', 'Core Foundation',    'Infrastructure, hosting, CI/CD, authentication (SSO, MFA), user & role management, audit logging, email notifications, backup system, admin console'],
            ['Phase 3', 'Financial Suite',    'General Ledger, Accounts Receivable, Accounts Payable, bank reconciliation, tax management, financial reporting (P&L, Balance Sheet, Cash Flow)'],
            ['Phase 4', 'CRM & Sales',        'Contact & account management, opportunity pipeline, quoting & estimates, order management, customer portal, sales reporting'],
            ['Phase 5', 'Inventory, POS & eCommerce', 'Product catalog, inventory tracking, purchase orders, warehouse management, POS terminal, eCommerce storefront, checkout, shipping & fulfillment'],
            ['Phase 6', 'Project & Job Mgmt','Job creation & tracking, milestones, resource & time tracking, job costing, art/proof approval workflow, production scheduling'],
            ['Phase 7', 'Reporting & Analytics', 'Executive dashboard, department dashboards, custom report builder, scheduled delivery, KPI tracking'],
            ['Phase 8', 'Compliance & Launch','PCI DSS audit, SOC 2 Type 2 controls, penetration testing, performance testing, disaster recovery drill, production launch, training, 30-day hypercare'],
        ],
        col_widths=[0.7, 1.5, 4.1],
    )

    # 6. Platform requirements
    add_heading(doc, '6.  Platform Requirements', 1)
    styled_table(doc,
        headers=['Requirement', 'Specification'],
        rows=[
            ['Hosting',               'Cloud-hosted, internet-accessible via browser'],
            ['Uptime Target',         '99.9% or greater'],
            ['Security Compliance',   'PCI DSS Level 1, SOC 2 Type 2 certification path'],
            ['Authentication',        'Multi-factor authentication (MFA), SSO support'],
            ['Data Backup',           'Automated daily backups, point-in-time recovery, off-site redundancy'],
            ['Disaster Recovery',     'Documented RTO ≤ 4 hours, RPO ≤ 1 hour'],
            ['Access Control',        'Role-based access control (RBAC), least-privilege model'],
            ['Audit Logging',         'Full user activity audit trail, tamper-evident, 3-year retention'],
            ['Data Encryption',       'TLS 1.2+ in transit, AES-256 at rest'],
            ['Notifications',         'Email notification system (transactional and alert-based)'],
            ['Browser Support',       'Chrome, Firefox, Safari, Edge (current versions)'],
            ['Responsive Design',     'Mobile-optimized web interface'],
        ],
        col_widths=[2.0, 4.3],
    )

    # 7. Assumptions
    add_heading(doc, '7.  Assumptions', 1)
    for i, item in enumerate([
        'Great Mountain West will provide a dedicated internal point of contact for the duration of the project.',
        'All stakeholder availability required for Phase 1 workshops will be provided within four (4) weeks of project start.',
        'Existing Sage 100 data export and migration is not included in this SOW unless explicitly added as a line item following Phase 1 assessment.',
        'Third-party software licenses (payment gateway, email provider, hosting) are not included and will be procured separately.',
        'Domain name, SSL certificates, and DNS management for the new platform are the responsibility of the client unless otherwise agreed.',
        'The client will designate a single feedback consolidator per department to prevent conflicting direction.',
        'Phase 2–8 pricing is not guaranteed and will be subject to change based on Phase 1 findings.',
    ], 1):
        add_body(doc, f'{i}.  {item}', size=10.5)

    # 8. Exclusions
    add_heading(doc, '8.  Exclusions', 1)
    for item in [
        'Data migration from Sage 100 to the new platform',
        'Legacy system decommissioning',
        'Hardware procurement',
        'Network or physical infrastructure',
        'Custom integrations not identified during Phase 1',
        'Mobile native applications (iOS/Android)',
        'Ongoing managed hosting or support (separate maintenance agreement)',
        'End-user training beyond the hypercare period',
    ]:
        add_bullet(doc, item)

    # 9. Out-of-scope rate
    add_heading(doc, '9.  Out-of-Scope & Change Order Rate', 1)
    add_body(doc, (
        'Work requested outside the defined scope of this SOW will be handled via a written Change Order. '
        'Change orders include a revised cost and timeline estimate and must be approved by both parties '
        'before work begins. Out-of-scope work is billed at $175.00 per hour.'
    ))

    # 10. IP
    add_heading(doc, '10.  Intellectual Property', 1)
    add_body(doc, (
        'Upon full payment for each phase, Great Mountain West shall own all custom code, designs, and '
        'documentation produced specifically for this project. Christopher Wall retains rights to any '
        'pre-existing frameworks, libraries, tools, or methodologies used in delivery.'
    ))

    # 11. Confidentiality
    add_heading(doc, '11.  Confidentiality', 1)
    add_body(doc, (
        'Both parties agree to treat all proprietary business information, technical specifications, and '
        'data shared during this engagement as confidential. A mutual Non-Disclosure Agreement (NDA) '
        'should be executed prior to discovery workshops.'
    ))

    # 12. Payment terms
    add_heading(doc, '12.  Payment Terms & Late Fees', 1)
    for item in [
        'Phase 1: Payment due in full prior to project commencement.',
        'Invoices for subsequent phases are due within fifteen (15) days of invoice date.',
        'Late payments are subject to a 1.5% monthly finance charge.',
        'Work may be paused on accounts with invoices past 30 days outstanding.',
    ]:
        add_bullet(doc, item)

    # 13. Change management
    add_heading(doc, '13.  Change Management', 1)
    add_body(doc, (
        'Changes to scope must be submitted in writing and approved by both parties via a Change Order '
        'before work begins. Change orders that increase scope will include a revised cost and timeline estimate.'
    ))

    # 14. Termination
    add_heading(doc, '14.  Termination', 1)
    add_body(doc, (
        'Either party may terminate this agreement with thirty (30) days written notice. Upon termination, '
        'Great Mountain West shall pay for all work completed to date. Phase 1 payment is '
        'non-refundable once work has commenced.'
    ))

    # 15. Limitation of liability
    add_heading(doc, '15.  Limitation of Liability', 1)
    add_body(doc, (
        "Christopher Wall's total liability shall not exceed the total fees paid under this SOW. "
        'Neither party shall be liable for indirect, incidental, or consequential damages.'
    ))

    # 16. Governing law
    add_heading(doc, '16.  Governing Law', 1)
    add_body(doc, 'This agreement shall be governed by the laws of the State of Utah.')

    add_page_break(doc)

    # 17. Signatures
    add_heading(doc, '17.  Signatures', 1)
    add_body(doc, (
        'By signing below, both parties agree to the terms set forth in this Statement of Work.'
    ))

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = 'Table Grid'

    def sig_block(cell, party, name, title):
        set_cell_bg(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(party)
        set_run_font(run, size=11, bold=True, color=DARK_NAVY)
        for label in [f'\n{name}', f'{title}', '\n\nSignature: ___________________________',
                      '\nPrinted Name: ___________________________',
                      '\nDate: ___________________________']:
            p2 = cell.add_paragraph(label)
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after = Pt(3)
            for run in p2.runs:
                set_run_font(run, size=10, color=MID_GRAY)

    sig_block(sig_table.rows[0].cells[0], 'Great Mountain West', '', '')
    sig_block(sig_table.rows[0].cells[1], 'Christopher Wall', 'Christopher Wall', 'Principal Developer')

    for cell in sig_table.rows[0].cells:
        cell.width = Inches(3.15)

    doc.add_paragraph()
    add_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'This document is confidential and intended solely for the use of '
        'Great Mountain West and Christopher Wall. Unauthorized distribution is prohibited.'
    )
    set_run_font(run, size=8.5, italic=True, color=MID_GRAY)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    # Default style tweaks
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    build_cover(doc)
    build_body(doc)

    out = r'C:\Projects\GMW\docs\sow\SOW_GreatMountainWest.docx'
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    main()
